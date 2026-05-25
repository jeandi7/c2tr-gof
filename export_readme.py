# export_readme.py — Convert README.md to README.docx
#
# Usage: python export_readme.py

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SOURCE = "README.md"
OUTPUT = "README.docx"

CODE_FONT = "Courier New"
CODE_BG   = RGBColor(0xF0, 0xF0, 0xF0)


# ── helpers ───────────────────────────────────────────────────────────────────

def _shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _apply_inline(run, text: str):
    """Apply bold/italic/inline-code markers inside a run's text."""
    # inline code `...`
    if text.startswith("`") and text.endswith("`") and len(text) > 2:
        run.text = text[1:-1]
        run.font.name = CODE_FONT
        run.font.size = Pt(9)
        return
    run.text = text


def _add_inline_paragraph(doc, text: str, style: str = "Normal") -> None:
    """Add a paragraph with basic inline markdown (bold, code)."""
    para = doc.add_paragraph(style=style)
    # split on **bold**, *italic*, `code`
    tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text)
    for token in tokens:
        run = para.add_run()
        if token.startswith("**") and token.endswith("**"):
            run.bold = True
            run.text = token[2:-2]
        elif token.startswith("`") and token.endswith("`"):
            run.text = token[1:-1]
            run.font.name = CODE_FONT
            run.font.size = Pt(9)
        elif token.startswith("*") and token.endswith("*"):
            run.italic = True
            run.text = token[1:-1]
        else:
            run.text = token


def _add_code_block(doc, lines: list[str]) -> None:
    for line in lines:
        para = doc.add_paragraph(style="Normal")
        para.paragraph_format.left_indent = Inches(0.3)
        run = para.add_run(line)
        run.font.name = CODE_FONT
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
        # light grey background via shading on paragraph
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F0F0F0")
        pPr.append(shd)
    doc.add_paragraph()  # spacing after block


def _add_table(doc, rows: list[str]) -> None:
    # rows = raw markdown table lines including the separator row
    data = []
    for row in rows:
        if re.match(r"^\|[-| :]+\|$", row.strip()):
            continue  # skip separator
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        data.append(cells)

    if not data:
        return

    n_cols = max(len(r) for r in data)
    table = doc.add_table(rows=len(data), cols=n_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(data):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", cell_text)
            for token in tokens:
                run = para.add_run()
                if token.startswith("**") and token.endswith("**"):
                    run.bold = True
                    run.text = token[2:-2]
                    if r_idx == 0:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif token.startswith("`") and token.endswith("`"):
                    run.text = token[1:-1]
                    run.font.name = CODE_FONT
                    run.font.size = Pt(8.5)
                else:
                    if r_idx == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.text = token
            if r_idx == 0:
                _shade_cell(cell, "2E74B5")
            elif r_idx % 2 == 0:
                _shade_cell(cell, "DEEAF1")

    doc.add_paragraph()  # spacing after table


# ── main converter ────────────────────────────────────────────────────────────

def convert(source: str, output: str) -> None:
    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    with open(source, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # ── code block ────────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            _add_code_block(doc, code_lines)
            i += 1
            continue

        # ── horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^-{3,}$", line.strip()):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── table ─────────────────────────────────────────────────────────────
        if line.strip().startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(lines[i].rstrip("\n"))
                i += 1
            _add_table(doc, table_rows)
            continue

        # ── headings ──────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            style = f"Heading {level}"
            _add_inline_paragraph(doc, text, style)
            i += 1
            continue

        # ── bullet list ───────────────────────────────────────────────────────
        if re.match(r"^[-*]\s+", line):
            text = re.sub(r"^[-*]\s+", "", line)
            _add_inline_paragraph(doc, text, "List Bullet")
            i += 1
            continue

        # ── blank line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── normal paragraph ──────────────────────────────────────────────────
        _add_inline_paragraph(doc, line)
        i += 1

    doc.save(output)
    print(f"Saved: {output}")


if __name__ == "__main__":
    convert(SOURCE, OUTPUT)
